"""
知识库文档处理模块
支持 PDF、Word 等格式的文档处理和分块
"""
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

import pdfplumber
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter

logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """文档元数据类"""
    filename: str
    file_path: str
    file_size: int  # bytes
    file_type: str
    page_count: Optional[int] = None
    chunk_count: int = 0
    processed_at: str = ""
    extra: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.extra is None:
            self.extra = {}


class DocumentProcessor:
    """文档处理器类"""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: Optional[List[str]] = None,
    ):
        """
        初始化文档处理器
        
        Args:
            chunk_size: 分块大小
            chunk_overlap: 分块重叠大小
            separators: 自定义分隔符列表
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", "。", "！", "？", "；", " ", ""]
        
        # 初始化分块器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=self.separators,
        )
        
        # 备用简单分块器
        self.simple_splitter = CharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separator="\n",
        )
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        处理单个文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            包含元数据和文档块的字典
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # 提取文本
        text = self._extract_text(file_path)
        
        # 获取元数据
        metadata = self._extract_metadata(file_path, text)
        
        # 分块处理
        chunks = self._split_text(text)
        metadata.chunk_count = len(chunks)
        
        # 创建 LangChain Document 对象
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    **metadata.__dict__,
                    "chunk_index": i,
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        logger.info(f"Processed document: {path.name}, pages={metadata.page_count}, chunks={len(chunks)}")
        
        return {
            "metadata": metadata,
            "documents": documents,
        }
    
    def process_multiple(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """
        批量处理多个文档
        
        Args:
            file_paths: 文档路径列表
            
        Returns:
            处理结果列表
        """
        results = []
        for file_path in file_paths:
            try:
                result = self.process_document(file_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process document {file_path}: {e}")
                results.append({
                    "error": str(e),
                    "file_path": file_path,
                })
        return results
    
    def _extract_text(self, file_path: str) -> str:
        """
        根据文件类型提取文本
        
        Args:
            file_path: 文件路径
            
        Returns:
            提取的文本内容
        """
        path = Path(file_path)
        file_type = path.suffix.lower()
        
        if file_type == ".pdf":
            return self._extract_from_pdf(file_path)
        elif file_type in [".docx", ".doc"]:
            return self._extract_from_word(file_path)
        elif file_type in [".txt", ".md", ".csv"]:
            return self._extract_from_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        从 PDF 提取文本
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            文本内容
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            logger.debug(f"Extracted text from PDF: {len(text)} characters")
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise
    
    def _extract_from_word(self, file_path: str) -> str:
        """
        从 Word 文档提取文本
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            文本内容
        """
        text = ""
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            logger.debug(f"Extracted text from Word: {len(text)} characters")
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from Word: {e}")
            raise
    
    def _extract_from_text(self, file_path: str) -> str:
        """
        从纯文本文件提取内容
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            文本内容
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            logger.debug(f"Extracted text from file: {len(text)} characters")
            return text
        except Exception as e:
            logger.error(f"Failed to read text file: {e}")
            raise
    
    def _extract_metadata(self, file_path: str, text: str) -> DocumentMetadata:
        """
        提取文档元数据
        
        Args:
            file_path: 文件路径
            text: 文档文本内容
            
        Returns:
            文档元数据对象
        """
        path = Path(file_path)
        
        # 获取文件大小
        file_size = path.stat().st_size if path.exists() else 0
        
        # 获取页数
        page_count = None
        file_type = path.suffix.lower()
        
        if file_type == ".pdf":
            try:
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
            except Exception:
                pass
        elif file_type in [".docx", ".doc"]:
            try:
                doc = DocxDocument(file_path)
                page_count = len(doc.sections) if doc.sections else None
            except Exception:
                pass
        
        return DocumentMetadata(
            filename=path.name,
            file_path=str(path.absolute()),
            file_size=file_size,
            file_type=file_type,
            page_count=page_count,
            processed_at=str(path.stat().st_mtime) if path.exists() else "",
        )
    
    def _split_text(self, text: str) -> List[str]:
        """
        分割文本为块
        
        Args:
            text: 原始文本
            
        Returns:
            文本块列表
        """
        try:
            chunks = self.text_splitter.split_text(text)
            
            # 如果分块数量过多，使用简单分块
            if len(chunks) > 1000:
                logger.warning(f"Too many chunks ({len(chunks)}), using simple splitter")
                chunks = self.simple_splitter.split_text(text)
            
            return chunks
        except Exception as e:
            logger.error(f"Failed to split text: {e}")
            # 降级为简单分块
            return self.simple_splitter.split_text(text)
    
    def create_custom_splitter(
        self,
        chunk_size: int,
        chunk_overlap: int,
        separators: List[str],
    ) -> RecursiveCharacterTextSplitter:
        """
        创建自定义分块器
        
        Args:
            chunk_size: 分块大小
            chunk_overlap: 分块重叠
            separators: 分隔符列表
            
        Returns:
            自定义分块器
        """
        return RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
        )


# 便捷函数
def process_document(file_path: str, **kwargs) -> Dict[str, Any]:
    """
    处理单个文档的便捷函数
    
    Args:
        file_path: 文档路径
        **kwargs: 其他参数
        
    Returns:
        处理结果
    """
    processor = DocumentProcessor(
        chunk_size=kwargs.get("chunk_size", 1000),
        chunk_overlap=kwargs.get("chunk_overlap", 200),
    )
    return processor.process_document(file_path)
